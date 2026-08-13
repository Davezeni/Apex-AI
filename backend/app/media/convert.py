"""Document conversion via a Markdown interchange hub.

Every conversion is `source → parse → markdown → write → target`, so N formats
need O(N) adapters. Tabular formats (CSV, XLSX) round-trip through Markdown
tables; prose formats (TXT, MD, HTML, DOCX) through Markdown text.
"""

from __future__ import annotations

import csv
import io
from pathlib import Path

# Imported lazily so optional formats degrade gracefully if a lib is missing.
try:
    from openpyxl import Workbook, load_workbook
except ImportError:  # pragma: no cover
    Workbook = load_workbook = None  # type: ignore

try:
    from docx import Document
except ImportError:  # pragma: no cover
    Document = None  # type: ignore


class ConversionError(Exception):
    pass


# --------------------------------------------------------------------------- #
# Parsers: format → Markdown
# --------------------------------------------------------------------------- #

def csv_to_markdown(text: str) -> str:
    rows = list(csv.reader(io.StringIO(text)))
    return rows_to_markdown(rows)


def xlsx_to_markdown(path: Path) -> str:
    if load_workbook is None:
        raise ConversionError("openpyxl is not installed")
    wb = load_workbook(path, read_only=True, data_only=True)
    ws = wb.active
    rows = [[("" if c is None else str(c)) for c in row] for row in ws.iter_rows(values_only=True)]
    return rows_to_markdown(rows)


def docx_to_markdown(path: Path) -> str:
    if Document is None:
        raise ConversionError("python-docx is not installed")
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        parts.append(rows_to_markdown(rows))
    return "\n\n".join(p for p in parts if p.strip())


def rows_to_markdown(rows: list[list[str]]) -> str:
    rows = [r for r in rows if any(c.strip() for c in r)]
    if not rows:
        return ""
    width = max(len(r) for r in rows)
    rows = [r + [""] * (width - len(r)) for r in rows]
    header = rows[0]
    sep = ["---"] * width
    body = rows[1:]
    lines = ["| " + " | ".join(header) + " |", "| " + " | ".join(sep) + " |"]
    lines += ["| " + " | ".join(r) + " |" for r in body]
    return "\n".join(lines)


# --------------------------------------------------------------------------- #
# Markdown → structured rows
# --------------------------------------------------------------------------- #

def markdown_table_to_rows(md: str) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in md.splitlines():
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.strip("|").split("|")]
        # Skip separator rows (e.g. | --- | --- |)
        if all(set(c) <= {"-", ":"} for c in cells if c):
            continue
        rows.append(cells)
    return rows


# --------------------------------------------------------------------------- #
# Writers: Markdown → format
# --------------------------------------------------------------------------- #

def markdown_to_csv(md: str) -> str:
    rows = markdown_table_to_rows(md)
    buf = io.StringIO()
    csv.writer(buf).writerows(rows)
    return buf.getvalue()


def markdown_to_xlsx(md: str, path: Path) -> None:
    if Workbook is None:
        raise ConversionError("openpyxl is not installed")
    rows = markdown_table_to_rows(md)
    wb = Workbook()
    ws = wb.active
    for row in rows:
        ws.append(row)
    wb.save(path)


def markdown_to_html(md: str) -> str:
    from html import escape

    rows = markdown_table_to_rows(md)
    if rows:
        parts = ["<table>"]
        for i, row in enumerate(rows):
            tag = "th" if i == 0 else "td"
            parts.append("<tr>" + "".join(f"<{tag}>{escape(c)}</{tag}>" for c in row) + "</tr>")
        parts.append("</table>")
        return "\n".join(parts)
    return "".join(f"<p>{escape(line)}</p>" for line in md.splitlines() if line.strip())


# --------------------------------------------------------------------------- #
# Public conversion API
# --------------------------------------------------------------------------- #

SUPPORTED = {".csv", ".xlsx", ".docx", ".txt", ".md", ".html", ".json"}


def to_markdown(path: Path) -> str:
    ext = path.suffix.lower()
    text = path.read_text(encoding="utf-8", errors="replace")
    if ext == ".csv":
        return csv_to_markdown(text)
    if ext == ".xlsx":
        return xlsx_to_markdown(path)
    if ext == ".docx":
        return docx_to_markdown(path)
    if ext in (".txt", ".md", ".html"):
        return text
    if ext == ".json":
        import json

        data = json.loads(text)
        if isinstance(data, list) and data and isinstance(data[0], dict):
            keys = list(data[0].keys())
            rows = [keys] + [[str(d.get(k, "")) for k in keys] for d in data]
            return rows_to_markdown(rows)
        return text
    raise ConversionError(f"unsupported source format: {ext}")


def from_markdown(md: str, target_ext: str, out_path: Path) -> None:
    ext = target_ext.lower()
    if ext == ".csv":
        out_path.write_text(markdown_to_csv(md), encoding="utf-8")
    elif ext == ".xlsx":
        markdown_to_xlsx(md, out_path)
    elif ext == ".html":
        out_path.write_text(markdown_to_html(md), encoding="utf-8")
    elif ext in (".md", ".txt"):
        out_path.write_text(md, encoding="utf-8")
    else:
        raise ConversionError(f"unsupported target format: {ext}")


def convert_file(src: Path, dst: Path) -> None:
    md = to_markdown(src)
    from_markdown(md, dst.suffix, dst)
